import PyPDF2
from docx import Document
import io
import streamlit as st

def extract_text_from_file(uploaded_file):
    """
    Extract text from uploaded PDF or DOCX file
    
    Args:
        uploaded_file: Streamlit uploaded file object
        
    Returns:
        str: Extracted text content
    """
    try:
        if uploaded_file.type == "application/pdf":
            return extract_text_from_pdf(uploaded_file)
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return extract_text_from_docx(uploaded_file)
        else:
            raise ValueError(f"Unsupported file type: {uploaded_file.type}")
    except Exception as e:
        st.error(f"Error extracting text from file: {str(e)}")
        return ""

def extract_text_from_pdf(uploaded_file):
    """
    Extract text from PDF file using PyPDF2
    
    Args:
        uploaded_file: Streamlit uploaded file object
        
    Returns:
        str: Extracted text content
    """
    try:
        # Create a file-like object from the uploaded file
        pdf_file = io.BytesIO(uploaded_file.getvalue())
        
        # Create PDF reader object
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        text = ""
        # Extract text from all pages
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text += page.extract_text() + "\n"
        
        return text.strip()
    
    except Exception as e:
        raise Exception(f"Error reading PDF file: {str(e)}")

def extract_text_from_docx(uploaded_file):
    """
    Extract text from DOCX file using python-docx
    
    Args:
        uploaded_file: Streamlit uploaded file object
        
    Returns:
        str: Extracted text content
    """
    try:
        # Create a file-like object from the uploaded file
        docx_file = io.BytesIO(uploaded_file.getvalue())
        
        # Load the document
        doc = Document(docx_file)
        
        text = ""
        # Extract text from all paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return text.strip()
    
    except Exception as e:
        raise Exception(f"Error reading DOCX file: {str(e)}")

def clean_text(text):
    """
    Clean and normalize extracted text
    
    Args:
        text: Raw extracted text
        
    Returns:
        str: Cleaned text
    """
    if not text:
        return ""
    
    # Remove excessive whitespace
    lines = [line.strip() for line in text.split('\n')]
    lines = [line for line in lines if line]  # Remove empty lines
    
    cleaned_text = '\n'.join(lines)
    
    # Replace multiple spaces with single space
    import re
    cleaned_text = re.sub(r'\s+', ' ', cleaned_text)
    
    return cleaned_text.strip()
